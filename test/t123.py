from pathlib import Path
from docxtpl import DocxTemplate  # pip install docxtpl
import docx
import pandas as pd
# import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base_dir = Path(__file__).parent
IQC_001_path = base_dir / "t123.txt"

T123_path = 'C:\\Users\\Zz\\PycharmProjects\\IQC\\template\\t123.docx'
output_path = 'C:\\Users\\Zz\\PycharmProjects\\IQC\\template\\t123-s.docx'

print(IQC_001_path)
print(T123_path)

IQC_path = T123_path.replace("123","321",1)
print(IQC_path)
doc = DocxTemplate(IQC_path)

IQC_path = str(IQC_001_path).replace("123","321",1)
print(IQC_path)
doc = DocxTemplate(IQC_path)

# document = docx.Document(T123_path)
# tables = document.tables
# table1 = tables[1]
# cells = table1._cells

# print(len(document.paragraphs))
# for i in document.paragraphs:  #遍历全部段落
#     print(i.text)

# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None
# delete_paragraph(document.paragraphs[0])
#
#
# for i in range(0, 6):
#     row = table1.rows[0]
#     row._element.getparent().remove(row._element)
#     # print(len(table1.rows))

# for cell in cells:
#     print(cell.text)

# for i in range(0, len(table1.rows)):
#     for j in range(0, len(table1.columns)):
#         print(i, j, table1.cell(i, j).text)

# document.save(output_path)
