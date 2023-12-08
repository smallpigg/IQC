from pathlib import Path
from docxtpl import DocxTemplate  # pip install docxtpl
import docx
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_empty_rows(table, n):
    # 增加对n和表格列数的比较
    for row in table.rows:
        X_cell = row.cells[n-1]
        if X_cell.text == 'nan' or X_cell.text == '':
            row._element.getparent().remove(row._element)

# 设置打印内容的显示宽度和长度
pd.set_option('display.max_rows', 500)
pd.set_option('display.max_columns', 100)
pd.set_option('display.width', 1000)

# 读取文件地址到变量中
base_dir = Path(__file__).parent
IQC_B_path = base_dir / "template/IQC-B.docx"
TB_IQC_B_path = base_dir / "template/TB-IQC-B.docx"
TB_IQC_000_path = base_dir / "template/TB-IQC-000.docx"
TZD_AAA_B_IQC_path = base_dir / "template/TZD-AAA-B-IQC.docx"
TZD_AAA_B_TB_IQC_path = base_dir / "template/TZD-AAA-B-TB-IQC.docx"
TZD_ABA_B_IQC_path = base_dir / "template/TZD-ABA-B-IQC.docx"
TZD_ABA_B_TB_IQC_path = base_dir / "template/TZD-ABA-B-TB-IQC.docx"
TZD_AAA_B_IQC_path_new = base_dir / "template/TZD-AAA-B-IQC-new.docx"
TZD_AAA_B_TB_IQC_path_new = base_dir / "template/TZD-AAA-B-TB-IQC-new.docx"
TZD_ABA_B_IQC_path_new = base_dir / "template/TZD-ABA-B-IQC-new.docx"
TZD_ABA_B_TB_IQC_path_new = base_dir / "template/TZD-ABA-B-TB-IQC-new.docx"

excel_path = base_dir / "result/result-output.xlsx"

# 文件保存地址
output_dir = base_dir / "output"
# Create output folder for the word documents
output_dir.mkdir(exist_ok=True)

# Convert Excel sheet to pandas dataframe
df = pd.read_excel(excel_path, sheet_name="Sheet1")

# 日期格式转换
# df["修改日期"] = pd.to_datetime(df["修改日期"]).dt.date
# df["申请日期"] = pd.to_datetime(df["申请日期"]).dt.date
# df["A版日期"] = pd.to_datetime(df["A版日期"]).dt.date
for col in df.columns:
    if "日期" in col:
        df[col] = pd.to_datetime(df[col]).dt.date


# 增加IQC文件编号
df["IQC文件编号"] = df["质量标准编号"]
for i in range(0, len(df)):
    str1 = df.loc[i, '质量标准编号']
    str1 = str1.replace("MAT","IQC",1)
    df.loc[i, 'IQC文件编号'] = str1

# 增加物料名称
# df["IQC物料名称"] = df["质量标准文件名称"]
# for i in range(0, len(df)):
#     str1 = df.loc[i, '质量标准文件名称']
#     str1 = str1.split(" ",2)
#     df.loc[i, 'IQC物料名称'] = str1[1]

# 增加IQC文件名称
df["IQC文件名称"] = df["物资名称"]
for i in range(0, len(df)):
    str1 = df.loc[i, 'IQC文件编号'][0:3] + df.loc[i, 'IQC文件编号'][8:11] + '_1_' + df.loc[i, 'IQC文件编号'] + '-' + df.loc[i, '版本1'] + '版 ' + df.loc[i, '物资名称'] + ' ' + '进货检验作业指导书.docx'
    df.loc[i, 'IQC文件名称'] = str1

# 增加IQC文件记录文件名称
df["IQC记录文件名称"] = df["物资名称"]
for i in range(0, len(df)):
    str1 = df.loc[i, 'IQC文件编号'][0:3] + df.loc[i, 'IQC文件编号'][8:11] + '_3_' + 'TB-' + df.loc[i, 'IQC文件编号'] + '-' + df.loc[i, '版本1'] + '版 ' + df.loc[
        i, '物资名称'] + ' ' + '进货检验记录.docx'
    df.loc[i, 'IQC记录文件名称'] = str1

# 增加IQC文件记录通知单文件名称
df["IQC文件通知单名称"] = df["物资名称"]
for i in range(0, len(df)):
    str1 = df.loc[i, 'IQC文件编号'][0:3] + df.loc[i, 'IQC文件编号'][8:11] + '_2_' + df.loc[i, 'IQC文件编号'] + '-' + df.loc[i, '版本1'] + '版 ' + df.loc[
        i, '物资名称'] + ' ' + '进货检验作业指导书文件记录更改通知单.docx'
    df.loc[i, 'IQC文件通知单名称'] = str1

# 增加TBIQC文件记录通知单文件名称
df["IQC记录文件通知单名称"] = df["物资名称"]
for i in range(0, len(df)):
    str1 = df.loc[i, 'IQC文件编号'][0:3] + df.loc[i, 'IQC文件编号'][8:11] + '_4_' + 'TB-' + df.loc[i, 'IQC文件编号'] + '-' + df.loc[i, '版本1'] + '版 ' + df.loc[
        i, '物资名称'] + ' ' + '进货检验记录文件记录更改通知单.docx'
    # str1 = df.loc[i, '质量标准文件名称']
    # str1 = str1.replace("MAT","IQC",1)
    # str2 = str1[12]
    # str3 = df.loc[i, 'IQC_TB版本']
    # str1 = str1.replace(str2, str3, 1)
    # str1 = "TB-"+str1.replace("质量标准", "进货检验记录文件记录更改通知单", 1)
    df.loc[i, 'IQC记录文件通知单名称'] = str1

# 读取记录单中填表需要的两个单元格内容
document = docx.Document(TB_IQC_000_path)
tables = document.tables
cell1 = tables[0].cell(1, 3)
cell2 = tables[0].cell(2, 3)

for record in df.to_dict(orient="records"):
    # 生成IQC文件
    # i = 2
    # while i <= 7:
    #     if str(record['检验项目'+str(i+1)]) == "nan":
    #         str1 = '00'+str(i-1)
    #         IQC_path = str(IQC_001_path).replace("001",str1,1)
    #         doc = DocxTemplate(IQC_path)
    #         doc.render(record)
    #         output_path = output_dir / f"{record['IQC文件名称']}"
    #         doc.save(output_path)
    #         document = docx.Document(output_path)
    #         tables = document.tables
    #         table1 = tables[-1]
    #         if str(record['IQC版本']) == "A":
    #             row = table1.rows[-1]
    #             row._element.getparent().remove(row._element)
    #             table1.cell(1, 1).text = str(record['申请日期'])
    #         document.save(output_path)
    #         break
    #     else:
    #         i += 1
    doc = DocxTemplate(IQC_B_path)
    doc.render(record)
    output_path = output_dir / f"{record['IQC文件名称']}"
    doc.save(output_path)

    checkbox_value = 1
    if checkbox_value:
        doc = docx.Document(output_path)
        for table in doc.tables:
            delete_empty_rows(table, 3)
        doc.save(output_path)
    


    # 生成记录单
    # 渲染表格内容
    doc = DocxTemplate(TB_IQC_B_path)
    doc.render(record)
    output_path = output_dir / f"{record['IQC记录文件名称']}"
    doc.save(output_path)

    # 填充表格
    document = docx.Document(output_path)
    tables = document.tables
    table1 = tables[0]
    table2 = tables[1]

    a = 0
    for i in range(1, 10):
        if record['项目'+str(i)] == "尺寸":
            table2.cell(2, 0).text = str(i) + '.'
            table2.cell(2, 0).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
            a += 1
        elif str(record['项目' + str(i)]) != "nan":
            list_string = [record['项目'+str(i)]]
            string_set = set(['材料', '材质', '产品包装', '单证资料', '规格型号', '合格证明', '认证资料', '产品描述'])
            row = table1.add_row()
            row.cells[0].text = str(i) + '.'
            row.cells[0].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
            row.cells[1].text = record['项目' + str(i)]
            row.cells[1].paragraphs[0].alignment = WD_TAB_ALIGNMENT.LEFT
            row.cells[2].text = record['接收标准' + str(i)]
            row.cells[2].paragraphs[0].alignment = WD_TAB_ALIGNMENT.LEFT
            if all([word in string_set for word in list_string]):
                run = table1.cell(len(table1.rows) - 1, 3).paragraphs[0].add_run(cell1.text)
            else:
                run = table1.cell(len(table1.rows) - 1, 3).paragraphs[0].add_run(cell2.text)
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            for cell in row.cells:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER
                set_cell_border(cell,
                                top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                left={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                right={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 4, "val": "single", "color": "#000000", "space": "0"})
            
            
            a += 1
        else:
            # row = table1.add_row()
            # row.cells[0].merge(row.cells[2])
            # row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # row.cells[0].text = "检验结果"
            
            # # table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            # # print(row)
            # cell_10 = row.cells[0]
            # paragraph = cell_10.paragraphs[0]
            # # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # run = paragraph.runs[0]
            # run.font.bold = True
            
            # for cell in row.cells:
            #     cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #     cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.LEFT
            #     set_cell_border(cell,
            #                     top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
            #                     bottom={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
            #                     left={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
            #                     right={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
            #                     insideH={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
            #                     end={"sz": 4, "val": "single", "color": "#000000", "space": "0"})
            break

    ChicunrowNum = 25
    #删除多余部分
    if table2.cell(2, 2).text == 'nan':
        for i in range(0, ChicunrowNum+3):
            row = table2.rows[0]
            row._element.getparent().remove(row._element)
        delete_paragraph(document.paragraphs[0])
        # row = table1.rows[len(table1.rows)-1]
        # row._element.getparent().remove(row._element)
    else:
        k = 0
        for i in range(3, ChicunrowNum+2):
            if table2.cell(i, 2).text == 'nan':
                k = i
                break
        for i in range(k, ChicunrowNum+2):
            row = table2.rows[k]
            row._element.getparent().remove(row._element)
    
    # 获取第一个表格
    table3 = document.tables[1]
    first_cell_text = ''
    # 获取第一列的第一行的文本
    # print(len(table3.rows))
    if len(table3.rows)>=1:
        first_cell_text = table3.cell(0, 0).text
    
    # print(first_cell_text)

    # 检查是否存在目标字符串
    str_to_delete = '检验结果'
    if first_cell_text == str_to_delete:
        # 如果存在，删除表格
        table3._element.clear()
    
    document.save(output_path)

    # 生成AAA两个通知单
    if record['质量标准编号'][:3] == 'AGA':
        if record['版本1'] == 'A':
            doc = DocxTemplate(TZD_AAA_B_IQC_path_new)
            doc.render(record)
            output_path = output_dir / f"{record['IQC文件通知单名称']}"
            doc.save(output_path)
        else:
            doc = DocxTemplate(TZD_AAA_B_IQC_path)
            doc.render(record)
            output_path = output_dir / f"{record['IQC文件通知单名称']}"
            doc.save(output_path)
        if record['版本1'] == 'A':
            doc = DocxTemplate(TZD_AAA_B_TB_IQC_path_new)
            doc.render(record)
            output_path = output_dir / f"{record['IQC记录文件通知单名称']}"
            doc.save(output_path)
        else:
            doc = DocxTemplate(TZD_AAA_B_TB_IQC_path)
            doc.render(record)
            output_path = output_dir / f"{record['IQC记录文件通知单名称']}"
            doc.save(output_path)
    # print('AAA', record['质量标准编号'][:3])
    # 生成ABA两个通知单
    else:
        if record['版本1'] == 'A':
            doc = DocxTemplate(TZD_ABA_B_IQC_path_new)
            doc.render(record)
            output_path = output_dir / f"{record['IQC文件通知单名称']}"
            doc.save(output_path)
        else:
            doc = DocxTemplate(TZD_ABA_B_IQC_path)
            doc.render(record)
            output_path = output_dir / f"{record['IQC文件通知单名称']}"
            doc.save(output_path)
        if record['版本1'] == 'A':
            doc = DocxTemplate(TZD_ABA_B_TB_IQC_path_new)
            doc.render(record)
            output_path = output_dir / f"{record['IQC记录文件通知单名称']}"
            doc.save(output_path)
        else:
            doc = DocxTemplate(TZD_ABA_B_TB_IQC_path)
            doc.render(record)
            output_path = output_dir / f"{record['IQC记录文件通知单名称']}"
            doc.save(output_path)
        # doc = DocxTemplate(TZD_ABA_B_IQC_path)
        # doc.render(record)
        # output_path = output_dir / f"{record['IQC文件通知单名称']}"
        # doc.save(output_path)
        # doc = DocxTemplate(TZD_ABA_B_TB_IQC_path)
        # doc.render(record)
        # output_path = output_dir / f"{record['IQC记录文件通知单名称']}"
        # doc.save(output_path)
        # print('ABA', record['质量标准编号'][:3])

    print(record['IQC文件编号']," is done!")