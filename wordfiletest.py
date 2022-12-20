import docx
import pandas as pd
import os
from pathlib import Path

base_dir = Path(__file__).parent
word_path = base_dir / "wordfiles/AAA-MAT-014-俯仰制动器固定盘质量标准（A版）.docx"


document = docx.Document(word_path)
tables = document.tables

cells = tables[0]._cells
cells1 = tables[len(tables) - 1]._cells

cells = cells + cells1

cells_text = [[cell.text for cell in cells]]
# cells_text1 = [[cell.text for cell in cells1]]

# cells_text = cells_text + cells_text1
print(cells_text)
#
# cells_text = cells_text.append(cells_text1)
# print(cells_text)

# result = []
# for element in cells_text:
#     result.append(element)
# for element in cells_text1:
#     result.append(element)
#
# print(result)


# df = pd.DataFrame(cells_text)
#
# row1 = [4,5,6,7]
# row2 = [16,17,18,19,20]
# row3 = [23,24,25,26,27]
#
# print(df[row1])
# print(df[row2])
# print(df[row3])


# for table in document.tables:
#     for row_index,row in enumerate(table.rows):
#         print(enumerate(table.rows))
#         for col_index,cell in enumerate(row.cells):
#             if row_index in range(1, 4) and col_index in range(2, 7):
#                 print('pos index is ({},{})'.format(row_index,col_index))
#                 print('cell text is {}'.format(cell.text))