import os
import pandas as pd
from docx import Document

def read_word_tables_to_df(filepath, table_numbers):
    doc = Document(filepath)
    
    # 获取正文中的表格
    body_tables = doc.tables
    
    # 获取页眉和页脚中的表格
    header_footer_tables = []
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            header_footer_tables.append(table)
        footer = section.footer
        for table in footer.tables:
            header_footer_tables.append(table)
    
    # 将正文、页眉、页脚中的表格合并为一个列表
    tables = body_tables + header_footer_tables

    data = []

    for n in table_numbers:
        if n > len(tables):
            raise ValueError(f"表格序号{n}超出了文档中的表格数量。")

        table = tables[n - 1]

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                column_name = f"{n}-{row_idx + 1}-{col_idx + 1}"
                data.append((column_name, cell.text))

    df = pd.DataFrame(data, columns=["column_name", "value"]).set_index("column_name").T

    return df

def process_batch_word_files(input_folder, output_file, table_numbers):
    dfs = []

    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            filepath = os.path.join(input_folder, filename)
            df = read_word_tables_to_df(filepath, table_numbers)
            dfs.append(df)

    result_df = pd.concat(dfs, ignore_index=True)
    result_df.to_excel(output_file, index=False)

# 示例用法
# input_folder = "word_files"  # 替换为包含Word文档的文件夹路径
# output_file = "output.xlsx"  # 替换为你想要保存结果的Excel文件路径
# table_numbers = [0, -1, 1]  # 替换为你想要读取的表格序号列表

# process_batch_word_files(input_folder, output_file, table_numbers)
