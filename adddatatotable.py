from pathlib import Path
from docxtpl import DocxTemplate  # pip install docxtpl
import docx
import pandas as pd
# import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

base_dir = Path(__file__).parent
output_dir = base_dir / "OUTPUT"
excel_path = base_dir / "wordtest.xlsx"


output_dir.mkdir(exist_ok=True)
df = pd.read_excel(excel_path, sheet_name="Sheet1")

# print(df)
TB_IQC_001_path = 'C:\\Users\\Zz\PycharmProjects\\IQC\\template\\TB-IQC-001.docx'
word_path = 'C:\\Users\\Zz\PycharmProjects\\IQC\\template\\wordtest.docx'

document = docx.Document(word_path)
tables = document.tables
table00 = tables[0]

cell1 = table00.cell(1, 3)
cell2 = table00.cell(2, 3)

print("111", cell1.text)
print("111", cell2.text)

for record in df.to_dict(orient="records"):
    document = docx.Document(TB_IQC_001_path)
    tables = document.tables
    table1 = tables[0]
    table2 = tables[1]

    a = 0
    for i in range(1, 7):
        if str(record['检验项目'+str(i+1)]) == "nan":
            #print("finished")
            row = table1.add_row()
            cell3 = row.cells[0]
            cell4 = row.cells[3]
            cell3.merge(cell4)
            row.cells[0].text = "备注："
            for cell in row.cells:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.LEFT
                set_cell_border(cell,
                                top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"})
            break
        elif record['检验项目'+str(i)] == "尺寸":
            table2.cell(2, 0).text = str(i) + '.'
            a += 1
        else:
            list_string = [record['检验项目'+str(i)]]
            string_set = set(['材料', '产品包装', '单证资料', '规格型号', '合格证明'])
            row = table1.add_row()
            row.cells[0].text = str(i) + '.'
            row.cells[1].text = record['检验项目' + str(i)]
            row.cells[2].text = record['检验项目' + str(i) + '接收标准']
            if all([word in string_set for word in list_string]):
                # row.cells[3].text = cell1.text
                run = table1.cell(len(table1.rows) - 1, 3).paragraphs[0].add_run(cell1.text)
            else:
                # row.cells[3].text = cell2.text
                run = table1.cell(len(table1.rows) - 1, 3).paragraphs[0].add_run(cell2.text)
            #run = table1.cell(len(table1.rows)-1, 3).paragraphs[0].add_run(cell1.text)
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            for cell in row.cells:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER
                set_cell_border(cell,
                                top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"})



            #print("one row added!")
            #print(len(tables))
            a += 1
    output_path = output_dir / f"{record['name']}"
    document.save(output_path)
    # print(a)




# document = docx.Document(word_path)
# tables = document.tables
# table1 = tables[0]
#
# cell1 = table1.cell(1, 3)
# cell2 = table1.cell(2, 3)
#
# row = table1.add_row()
#
# run = table1.cell(3, 3).paragraphs[0].add_run(cell1.text)
# run.font.name = u'宋体'
# run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
# for cell in row.cells:
#     cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER
#     set_cell_border(cell,
#                     top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
#                     bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
#                     left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
#                     right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
#                     insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
#                     end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"})
#
# print(row)
# print(cell2.text)
#
# # row = table1.rows[-1]
# # str1 = row.cells[0].text
# # print(str1)
#
# output_path = output_dir / "output.docx"
# document.save(output_path)





# for record in df.to_dict(orient="records"):
#     document = docx.Document(word_path)
#     tables = document.tables
#     table1 = tables[0]
#
#     a = 0
#     for i in range(1, 7):
#         if str(record['检验项目'+str(i+1)]) == "nan":
#             #print("finished")
#             break
#         else:
#             table1.rows[1].add_row()
#             #print("one row added!")
#             #print(len(tables))
#             a += 1
#             output_path = output_dir / f"{record['name']}"
#             document.save(output_path)
#     print(a)
#
#


    #         str1 = '00' + str(i - 1)
    #         IQC_path = IQC_001_path.replace("001",str1,1)
    #         doc = DocxTemplate(IQC_path)
    #         doc.render(record)
    #         output_path = output_dir / f"{record['IQC文件名称']}"
    #         doc.save(output_path)
    #         a += 1
    #
    #
    # print(record)
#     i = 2
#     while i <= 7:
#         if str(record['检验项目'+str(i+1)]) == "nan":
#             str1 = '00'+str(i-1)
#             IQC_path = IQC_001_path.replace("001",str1,1)
#             doc = DocxTemplate(IQC_path)
#             doc.render(record)
#             output_path = output_dir / f"{record['IQC文件名称']}"
#             doc.save(output_path)
#             break
#         else:
#             i += 1